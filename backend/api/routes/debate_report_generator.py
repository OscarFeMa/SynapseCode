# ============================================================================
# GENERADOR DE INFORMES PROFESIONALES - ENFOQUE HIBRIDO
# Datos/Tablas: Programatico | Narrativa: LLM | Verificacion: Automatica
# ============================================================================

from datetime import datetime

from fastapi import HTTPException, Response


class StrategyData:
    """Estructura para datos extraidos de estrategias"""

    def __init__(self, name: str):
        self.name = name
        self.capital_raw: str = ""
        self.capital_value: float = 0
        self.retorno_raw: str = ""
        self.riesgo_raw: str = ""
        self.riesgo_level: str = "medio"
        self.plazo: str = ""
        self.mencionado_por: list[str] = []
        self.descripcion: str = ""
        self.pasos: list[str] = []
        self.riesgos_mencionados: list[str] = []


def _extract_number(text: str) -> float:
    """Extrae el primer numero de un texto (maneja $, comas, puntos)"""
    import re

    # Buscar patron con $ y numero (puede tener comas)
    match = re.search(r"[\$€]\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        return float(match.group(1).replace(",", ""))
    # Buscar numero solo
    match = re.search(r"(\d{1,3}(?:,\d{3})*(?:\.\d+)?)", text)
    if match:
        return float(match.group(1).replace(",", ""))
    # Buscar porcentaje
    match = re.search(r"(\d+(?:\.\d+)?)\s*(%|por ciento)", text.lower())
    if match:
        return float(match.group(1))
    return 0


def _extract_strategies_hybrid(completed_turns: list) -> list[StrategyData]:
    """
    Extrae estrategias con datos concretos usando parsing programatico.
    Busca patrones estructurados en el contenido de los turnos.
    """
    strategies: dict[str, StrategyData] = {}

    for turn in completed_turns:
        content = turn.get("response_received", "")
        model = turn.get("model", "")
        lines = content.split("\n")

        current_strategy = None

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            lower = stripped.lower()

            # Detectar inicio de estrategia (formato: **Estrategia X: Nombre**)
            is_strategy = False
            name = ""

            # Formato 1: **Estrategia 1: Trading de Alta Frecuencia**
            if stripped.startswith("**Estrategia") or stripped.startswith("**estrategia"):
                is_strategy = True
                # Extraer nombre despues del primer ":"
                if ":" in stripped:
                    name = stripped.split(":", 1)[1].replace("**", "").strip()
                else:
                    name = stripped.replace("**", "").replace("Estrategia", "").replace("estrategia", "").strip()

            # Formato 2: # Estrategia 1: Nombre
            elif stripped.startswith("# Estrategia") or stripped.startswith("# estrategia"):
                is_strategy = True
                if ":" in stripped:
                    name = stripped.split(":", 1)[1].replace("#", "").strip()
                else:
                    name = stripped.replace("#", "").replace("Estrategia", "").replace("estrategia", "").strip()

            if is_strategy and name and len(name) > 3 and len(name) < 100:
                name = name.title()
                if name not in strategies:
                    strategies[name] = StrategyData(name)
                current_strategy = strategies[name]
                if model not in current_strategy.mencionado_por:
                    current_strategy.mencionado_por.append(model)
                continue

            if current_strategy:
                # Extraer capital (formato: "* Capital inicial estimado: $10,000")
                if "capital" in lower or "inversion inicial" in lower:
                    current_strategy.capital_raw = stripped
                    val = _extract_number(stripped)
                    if val > 0:
                        current_strategy.capital_value = val

                # Extraer retorno
                elif "retorno" in lower or "rentabilidad" in lower or "ganancia" in lower or "roi" in lower:
                    current_strategy.retorno_raw = stripped

                # Extraer riesgo
                elif "riesgo" in lower or "peligro" in lower:
                    current_strategy.riesgo_raw = stripped
                    if "alto" in lower or "elevado" in lower:
                        current_strategy.riesgo_level = "alto"
                    elif "bajo" in lower:
                        current_strategy.riesgo_level = "bajo"
                    elif "medio" in lower or "moderado" in lower:
                        current_strategy.riesgo_level = "medio"

                # Extraer plazo
                elif "plazo" in lower or "semana" in lower or "mes" in lower:
                    if not current_strategy.plazo:
                        current_strategy.plazo = stripped[:100]

                # Extraer pasos accionables (solo si no es una linea de metadata)
                elif (
                    stripped.startswith(("+", "-", "•"))
                    and "capital" not in lower
                    and "retorno" not in lower
                    and "riesgo" not in lower
                ):
                    if len(stripped) > 15:
                        current_strategy.pasos.append(stripped[:200])

                # Extraer riesgos mencionados
                elif any(kw in lower for kw in ["volatilidad", "perdida", "perder"]):
                    current_strategy.riesgos_mencionados.append(stripped[:150])

                # Descripcion general
                elif (
                    len(stripped) > 30
                    and not stripped.startswith("#")
                    and not stripped.startswith("*")
                    and not current_strategy.descripcion
                ):
                    current_strategy.descripcion = stripped[:300]

    return sorted(
        strategies.values(),
        key=lambda s: s.capital_value if s.capital_value > 0 else 999999,
    )


def _build_programmatic_sections(strategies: list[StrategyData], completed_turns: list, failed_summary: str) -> str:
    """Genera las secciones de datos directamente sin LLM"""

    unique_models = set(t.get("model", "") for t in completed_turns)
    unique_roles = set(t.get("agent_role", "") for t in completed_turns)

    total_tokens = sum(t.get("tokens_out", 0) for t in completed_turns)
    total_latency = sum(t.get("latency_ms", 0) for t in completed_turns)
    avg_latency = total_latency / max(len(completed_turns), 1)

    sections = []

    # 1. Tabla de estrategias (datos exactos)
    if strategies:
        table_lines = ["| # | Estrategia | Capital | Retorno | Riesgo | Plazo | Mencionada por |"]
        table_lines.append("|---|-----------|---------|---------|--------|-------|----------------|")
        for i, s in enumerate(strategies, 1):
            capital = f"${s.capital_value:,.0f}" if s.capital_value > 0 else "No especificado"
            retorno = s.retorno_raw[:50] if s.retorno_raw else "No especificado"
            riesgo = s.riesgo_level.capitalize()
            plazo = s.plazo[:30] if s.plazo else "No especificado"
            modelos = ", ".join(s.mencionado_por[:3])
            table_lines.append(f"| {i} | {s.name} | {capital} | {retorno} | {riesgo} | {plazo} | {modelos} |")
        sections.append("\n".join(table_lines))

    # 2. Ranking por viabilidad
    if strategies:
        ranked = sorted(
            strategies,
            key=lambda s: (
                s.capital_value if s.capital_value > 0 else 999999,
                {"bajo": 0, "medio": 1, "alto": 2}.get(s.riesgo_level, 1),
            ),
        )
        ranking_lines = ["\n## Ranking de Viabilidad"]
        ranking_lines.append("\nOrdenadas por menor capital inicial y menor riesgo:")
        for i, s in enumerate(ranked, 1):
            capital = f"${s.capital_value:,.0f}" if s.capital_value > 0 else "N/A"
            ranking_lines.append(f"{i}. **{s.name}** - Capital: {capital} | Riesgo: {s.riesgo_level}")
            if s.pasos:
                ranking_lines.append(f"   - Primer paso: {s.pasos[0][:100]}")
        sections.append("\n".join(ranking_lines))

    # 3. Metricas del debate
    metrics = [
        "\n## Metricas del Debate",
        f"- **Intervenciones exitosas:** {len(completed_turns)}",
        f"- **Modelos participantes:** {len(unique_models)} ({', '.join(sorted(unique_models))})",
        f"- **Roles cubiertos:** {', '.join(sorted(unique_roles))}",
        f"- **Estrategias identificadas:** {len(strategies)}",
        f"- **Tokens generados:** {total_tokens:,}",
        f"- **Tiempo total:** {total_latency / 1000:.0f}s ({avg_latency / 1000:.1f}s promedio por turno)",
    ]
    if failed_summary:
        metrics.append(f"- **Nota:** {failed_summary.strip()}")
    sections.append("\n".join(metrics))

    return "\n\n".join(sections)


def _build_narrative_prompt(topic: str, strategies: list[StrategyData], completed_turns: list) -> str:
    """Construye prompt SOLO para la narrativa (resumen + analisis)"""

    analyst_content = ""
    for t in completed_turns:
        if t.get("agent_role") in ("analyst", "critic"):
            content = t.get("response_received", "")
            model = t.get("model", "")
            role = t.get("agent_role", "")
            analyst_content += f"\n[{model} - {role}]: {content[:400]}\n"

    synthesis_content = ""
    for t in completed_turns:
        if t.get("agent_role") in ("synthesizer", "moderator", "refiner"):
            content = t.get("response_received", "")
            model = t.get("model", "")
            synthesis_content += f"\n[{model}]: {content[:400]}\n"

    strategy_list = "\n".join(
        f"- {s.name} (Capital: ${s.capital_value:,.0f} if s.capital_value > 0 else 'N/A', Riesgo: {s.riesgo_level})"
        for s in strategies
    )

    prompt = f"""Eres un redactor financiero profesional con 15 anos de experiencia. Escribe SOLO las secciones narrativas de un informe.

TEMA: {topic}

ESTRATEGIAS IDENTIFICADAS (datos ya estan en tablas separadas):
{strategy_list}

ANALISIS DE EXPERTOS:
{analyst_content[:3000]}

SINTESIS DEL DEBATE:
{synthesis_content[:2000]}

ESCRIBE SOLO ESTAS 3 SECCIONES (sin tablas, sin metricas, solo texto fluido):

### Resumen Ejecutivo
2-3 parrafos que capturen la esencia del debate y la conclusion principal.

### Analisis del Debate
3-4 parrafos narrando como evoluciono el debate: que se propuso primero, que se cuestiono, donde hubo consenso. Usa transiciones naturales.

### Conclusion y Recomendacion Final
2 parrafos con el veredicto final y cual es la mejor estrategia global.

REGLAS:
- Espanol profesional
- Sin lenguaje tecnico de IA
- Usa las estrategias de la lista como referencia
- NO inventes cifras (las tablas ya las tienen)
- NO repitas las tablas o metricas (ya se generan por separado)

Genera solo el texto narrativo:"""

    return prompt


async def _generate_narrative_with_llm(prompt: str) -> str:
    """Genera solo la narrativa usando qwen2.5:7b en Master"""
    from backend.adapters.ollama import OllamaClient

    client = OllamaClient()
    full_response = []

    async for token in client.generate(
        model="qwen2.5:7b",
        prompt=prompt,
        system="Eres un redactor financiero profesional. Escribes texto narrativo fluido y natural. Nunca inventas cifras ni datos.",
        stream=True,
        options={"num_ctx": 8192, "temperature": 0.3, "num_predict": 3000},
    ):
        full_response.append(token)

    return "".join(full_response)


def _verify_report_figures(report: str, strategies: list[StrategyData]) -> list[str]:
    """Verifica que las cifras del informe coincidan con los datos reales"""
    import re

    warnings = []
    found_numbers = re.findall(r"[\$€]?\s*([\d,]+(?:\.\d+)?)", report)
    found_numbers = [n.replace(",", "") for n in found_numbers]

    for s in strategies:
        if s.capital_value > 0:
            capital_str = str(int(s.capital_value))
            if capital_str not in found_numbers:
                found_close = any(
                    abs(float(n) - s.capital_value) / max(s.capital_value, 1) < 0.1 for n in found_numbers if n
                )
                if not found_close:
                    warnings.append(
                        f"Cifra de capital para '{s.name}' (${s.capital_value:,.0f}) no encontrada en el informe"
                    )

    return warnings


async def generate_professional_report(session_id: str, debate_controller):
    """
    Genera un informe profesional con enfoque hibrido:
    - Datos/Tablas: Extraidos programaticamente (100% exactos)
    - Narrativa: Generada por LLM (qwen2.5:7b en Master)
    - Verificacion: Automatica de cifras
    Funciona desde memoria o base de datos (cualquier debate guardado).
    """

    # 1. Obtener datos del debate (memoria o BD)
    session = debate_controller.get_session(session_id)

    if session:
        topic = session.topic
        all_turns = []
        for t in session.turns:
            if hasattr(t, "turn_number"):
                all_turns.append(
                    {
                        "turn_number": t.turn_number,
                        "agent_name": t.agent.name,
                        "agent_role": t.agent.role.value if hasattr(t.agent.role, "value") else str(t.agent.role),
                        "model": t.agent.model,
                        "response_received": t.response_received,
                        "status": t.status,
                        "tokens_out": t.tokens_out,
                        "latency_ms": t.latency_ms,
                    }
                )
            else:
                all_turns.append(dict(t))
    else:
        debate_data = await debate_controller.get_debate_from_db(session_id)
        if not debate_data:
            raise HTTPException(status_code=404, detail="Debate not found in memory or database")

        topic = debate_data.get("topic", "Sin tema")
        all_turns = debate_data.get("turns", [])

    # 2. Separar turnos exitosos de fallidos
    completed_turns = [t for t in all_turns if str(t.get("status", "")).startswith("completed")]
    failed_turns = [t for t in all_turns if str(t.get("status", "")).startswith("failed")]

    # 3. Resumen de fallos
    failed_summary = ""
    if failed_turns:
        failed_models = {}
        for ft in failed_turns:
            model = ft.get("model", "unknown")
            failed_models[model] = failed_models.get(model, 0) + 1
        failed_summary = "Fallos tecnicos: " + ", ".join(f"{m} ({c} turnos)" for m, c in failed_models.items())

    # 4. Extraer estrategias programaticamente
    strategies = _extract_strategies_hybrid(completed_turns)

    # 5. Generar secciones programaticas (tablas, ranking, metricas)
    programmatic_sections = _build_programmatic_sections(strategies, completed_turns, failed_summary)

    # 6. Generar narrativa con LLM
    narrative_prompt = _build_narrative_prompt(topic, strategies, completed_turns)

    try:
        narrative = await _generate_narrative_with_llm(narrative_prompt)
    except Exception as e:
        narrative = f"""### Resumen Ejecutivo
No se pudo generar la narrativa automatica. Error: {str(e)}

### Analisis del Debate
Ver las tablas de datos a continuacion.

### Conclusion
Ver las tablas de datos a continuacion."""

    # 7. Ensamblar informe final
    clean_topic = topic.split("\n")[0][:100]

    report = f"""# Informe Synapse Council

## {clean_topic}

---

{narrative}

---

{programmatic_sections}

---

*Informe generado automaticamente por Synapse Council v2.8*
*Metodo: Enfoque hibrido (datos programaticos + narrativa LLM)*
*Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}*
"""

    # 8. Verificar cifras
    _verify_report_figures(report, strategies)

    # 9. Retornar como Markdown
    return Response(
        content=report,
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename=report_{session_id[:8]}.md"},
    )


async def _build_report_content(session_id: str, debate_controller):
    """Construye el contenido del informe (reutilizable para DOCX/PDF)"""
    session = debate_controller.get_session(session_id)

    if session:
        topic = session.topic
        all_turns = []
        for t in session.turns:
            if hasattr(t, "turn_number"):
                all_turns.append(
                    {
                        "turn_number": t.turn_number,
                        "agent_name": t.agent.name,
                        "agent_role": t.agent.role.value if hasattr(t.agent.role, "value") else str(t.agent.role),
                        "model": t.agent.model,
                        "response_received": t.response_received,
                        "status": t.status,
                        "tokens_out": t.tokens_out,
                        "latency_ms": t.latency_ms,
                    }
                )
            else:
                all_turns.append(dict(t))
    else:
        debate_data = await debate_controller.get_debate_from_db(session_id)
        if not debate_data:
            return None, None, None, None, None

        topic = debate_data.get("topic", "Sin tema")
        all_turns = debate_data.get("turns", [])

    completed_turns = [t for t in all_turns if str(t.get("status", "")).startswith("completed")]
    failed_turns = [t for t in all_turns if str(t.get("status", "")).startswith("failed")]

    failed_summary = ""
    if failed_turns:
        failed_models = {}
        for ft in failed_turns:
            model = ft.get("model", "unknown")
            failed_models[model] = failed_models.get(model, 0) + 1
        failed_summary = "Fallos tecnicos: " + ", ".join(f"{m} ({c} turnos)" for m, c in failed_models.items())

    strategies = _extract_strategies_hybrid(completed_turns)
    programmatic_sections = _build_programmatic_sections(strategies, completed_turns, failed_summary)

    narrative_prompt = _build_narrative_prompt(topic, strategies, completed_turns)

    try:
        narrative = await _generate_narrative_with_llm(narrative_prompt)
    except Exception as e:
        narrative = f"""### Resumen Ejecutivo
No se pudo generar la narrativa automatica. Error: {str(e)}

### Analisis del Debate
Ver las tablas de datos a continuacion.

### Conclusion
Ver las tablas de datos a continuacion."""

    clean_topic = topic.split("\n")[0][:100]

    return topic, clean_topic, narrative, strategies, programmatic_sections


async def generate_report_as_docx(session_id: str, debate_controller):
    """Genera informe profesional como documento Word (.docx)"""
    from io import BytesIO

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    topic, clean_topic, narrative, strategies, programmatic_sections = await _build_report_content(
        session_id, debate_controller
    )

    if topic is None:
        raise HTTPException(status_code=404, detail="Debate not found")

    doc = Document()

    # Configurar estilos
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # PORTADA
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_heading("Synapse Council", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        run.font.size = Pt(36)

    subtitle = doc.add_heading("Informe de Analisis Estrategico", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.font.size = Pt(18)

    doc.add_paragraph("")

    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic_para.add_run(clean_topic)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.bold = True

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").font.size = Pt(10)
    meta.add_run("Metodo: Enfoque hibrido (datos programaticos + narrativa LLM)\n").font.size = Pt(10)
    meta.add_run("Synapse Council v2.8").font.size = Pt(10)

    doc.add_page_break()

    # NARRATIVA (Resumen, Analisis, Conclusion)
    doc.add_heading("Resumen Ejecutivo", level=1)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # Parsear narrativa para extraer secciones
    narrative_lines = narrative.split("\n")

    for line in narrative_lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            section_title = stripped.replace("### ", "")
            doc.add_heading(section_title, level=2)
            for run in doc.paragraphs[-1].runs:
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        elif stripped.startswith("## "):
            section_title = stripped.replace("## ", "")
            doc.add_heading(section_title, level=1)
            for run in doc.paragraphs[-1].runs:
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.replace("**", ""))
            run.bold = True
            run.font.size = Pt(12)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # TABLA DE ESTRATEGIAS
    doc.add_heading("Estrategias Identificadas", level=1)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    if strategies:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Shading Accent 1"

        # Headers
        headers = ["#", "Estrategia", "Capital", "Riesgo", "Mencionada por"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for i, s in enumerate(strategies, 1):
            row = table.add_row()
            capital = f"${s.capital_value:,.0f}" if s.capital_value > 0 else "N/A"
            modelos = ", ".join(s.mencionado_por[:2])
            row.cells[0].text = str(i)
            row.cells[1].text = s.name
            row.cells[2].text = capital
            row.cells[3].text = s.riesgo_level.capitalize()
            row.cells[4].text = modelos

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_page_break()

    # METRICAS
    doc.add_heading("Metricas del Debate", level=1)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    unique_models = set()
    total_tokens = 0
    total_latency = 0

    # Re-obtener datos para metricas
    session = debate_controller.get_session(session_id)
    if session:
        for t in session.turns:
            if hasattr(t, "turn_number") and t.status.startswith("completed"):
                unique_models.add(t.agent.model)
                total_tokens += t.tokens_out
                total_latency += t.latency_ms

    metrics = [
        (
            "Intervenciones exitosas",
            str(
                len(
                    [
                        t
                        for t in (session.turns if session else [])
                        if hasattr(t, "turn_number") and t.status.startswith("completed")
                    ]
                )
            ),
        ),
        ("Modelos participantes", f"{len(unique_models)}"),
        ("Estrategias identificadas", str(len(strategies))),
        ("Tokens generados", f"{total_tokens:,}"),
        ("Tiempo total", f"{total_latency / 1000:.0f}s"),
    ]

    for label, value in metrics:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    # FOOTER
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Informe generado automaticamente por Synapse Council v2.8").font.size = Pt(9)
    footer.add_run("\nMetodo: Enfoque hibrido (datos programaticos + narrativa LLM)").font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=report_{session_id[:8]}.docx"},
    )


async def generate_report_as_pdf(session_id: str, debate_controller):
    """Genera informe profesional como PDF usando xhtml2pdf"""
    import html as html_module
    import traceback
    from io import BytesIO

    from xhtml2pdf import pisa

    try:
        topic, clean_topic, narrative, strategies, programmatic_sections = await _build_report_content(
            session_id, debate_controller
        )

        if topic is None:
            raise HTTPException(status_code=404, detail="Debate not found")

        # Escapar contenido para HTML
        safe_topic = html_module.escape(clean_topic)
        safe_narrative = html_module.escape(narrative).replace("\n", "<br>")

        # Construir HTML para PDF
        html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
    @page {{
        size: letter;
        margin: 2cm;
        @top-right {{
            content: "Synapse Council v2.8";
            font-size: 8pt;
            color: #64748B;
        }}
        @bottom-center {{
            content: "Pagina " counter(page) " de " counter(pages);
            font-size: 8pt;
            color: #64748B;
        }}
    }}
    body {{
        font-family: 'Segoe UI', Calibri, Arial, sans-serif;
        color: #333;
        line-height: 1.6;
        font-size: 11pt;
    }}
    h1 {{
        color: #2563EB;
        border-bottom: 2px solid #E2E8F0;
        padding-bottom: 8px;
        font-size: 20pt;
        page-break-before: always;
    }}
    h1:first-of-type {{
        page-break-before: avoid;
    }}
    h2 {{
        color: #475569;
        font-size: 14pt;
        margin-top: 20px;
    }}
    h3 {{
        color: #64748B;
        font-size: 12pt;
    }}
    .cover {{
        text-align: center;
        padding-top: 100px;
    }}
    .cover h1 {{
        color: #2563EB;
        font-size: 36pt;
        border: none;
    }}
    .cover h2 {{
        color: #64748B;
        font-size: 18pt;
    }}
    .cover .topic {{
        font-size: 14pt;
        font-weight: bold;
        color: #1E293B;
        margin: 30px 0;
    }}
    .cover .meta {{
        font-size: 10pt;
        color: #64748B;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 20px 0;
    }}
    th {{
        background-color: #2563EB;
        color: white;
        padding: 10px;
        text-align: left;
        font-size: 10pt;
    }}
    td {{
        padding: 8px 10px;
        border-bottom: 1px solid #E2E8F0;
        font-size: 9pt;
    }}
    tr:nth-child(even) {{
        background-color: #F8FAFC;
    }}
    .metric {{
        margin: 10px 0;
    }}
    .metric strong {{
        color: #2563EB;
    }}
    .footer {{
        text-align: center;
        font-size: 9pt;
        color: #64748B;
        margin-top: 40px;
        border-top: 1px solid #E2E8F0;
        padding-top: 20px;
    }}
</style>
</head>
<body>

<div class="cover">
    <h1>Synapse Council</h1>
    <h2>Informe de Analisis Estrategico</h2>
    <div class="topic">{safe_topic}</div>
    <div class="meta">
        Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}<br>
        Metodo: Enfoque hibrido (datos programaticos + narrativa LLM)<br>
        Synapse Council v2.8
    </div>
</div>

<h1>Resumen Ejecutivo</h1>
{safe_narrative}

<h1>Estrategias Identificadas</h1>
"""

        if strategies:
            html_content += """<table>
<thead>
<tr>
    <th>#</th>
    <th>Estrategia</th>
    <th>Capital</th>
    <th>Riesgo</th>
    <th>Mencionada por</th>
</tr>
</thead>
<tbody>
"""
            for i, s in enumerate(strategies, 1):
                capital = f"${s.capital_value:,.0f}" if s.capital_value > 0 else "N/A"
                modelos = html_module.escape(", ".join(s.mencionado_por[:2]))
                name = html_module.escape(s.name)
                html_content += f"""<tr>
    <td>{i}</td>
    <td>{name}</td>
    <td>{capital}</td>
    <td>{s.riesgo_level.capitalize()}</td>
    <td>{modelos}</td>
</tr>
"""
            html_content += "</tbody></table>"

        # Metricas
        html_content += """<h1>Metricas del Debate</h1>"""

        unique_models = set()
        total_tokens = 0
        total_latency = 0
        completed_count = 0

        session = debate_controller.get_session(session_id)
        if session:
            for t in session.turns:
                if hasattr(t, "turn_number") and t.status.startswith("completed"):
                    unique_models.add(t.agent.model)
                    total_tokens += t.tokens_out
                    total_latency += t.latency_ms
                    completed_count += 1

        metrics = [
            ("Intervenciones exitosas", str(completed_count)),
            ("Modelos participantes", str(len(unique_models))),
            ("Estrategias identificadas", str(len(strategies))),
            ("Tokens generados", f"{total_tokens:,}"),
            ("Tiempo total", f"{total_latency / 1000:.0f}s"),
        ]

        for label, value in metrics:
            html_content += f'<div class="metric"><strong>{label}:</strong> {value}</div>'

        html_content += """
<div class="footer">
    Informe generado automaticamente por Synapse Council v2.8<br>
    Metodo: Enfoque hibrido (datos programaticos + narrativa LLM)
</div>

</body>
</html>"""

        buffer = BytesIO()
        pisa.CreatePDF(html_content, dest=buffer)
        buffer.seek(0)

        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=report_{session_id[:8]}.pdf"},
        )
    except Exception as e:
        error_detail = f"PDF generation error: {str(e)}\n{traceback.format_exc()}"
        print(error_detail)
        raise HTTPException(status_code=500, detail=error_detail)
